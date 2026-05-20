"""
TEST: python-docx fill Word template → upload Google Drive → export PDF

Cách chạy:
  python3 test-report-gen.py

Yêu cầu:
  - FIREBASE_SERVICE_ACCOUNT env variable (JSON string)
  - Drive API enabled trong Google Cloud Console (cùng project với Firebase)
  - pip install python-docx google-auth google-api-python-client
"""

import os, io, json, shutil, warnings
warnings.filterwarnings('ignore')

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# ── Dữ liệu test ────────────────────────────────────────────────────
TEST_DATA = {
    "check_tat_ca_nd": False,
    "check_co_mau_phat_hien": True,
    "r2": "0.9987",
    "nguoi_phan_tich": "Ong Thanh Dat",
    "ngay_phan_tich": "19/05/2026",
    "nguoi_tham_tra": "Nguyen Hoang Dao",
    "ngay_tham_tra": "19/05/2026",
    "samples": [
        {"lo_so": "1",  "ma_so_mau": "M01-2026-TFL",  "kq": None,    "ghi_chu": ""},
        {"lo_so": "2",  "ma_so_mau": "M02-2026-TFL",  "kq": "0.023", "ghi_chu": ""},
        {"lo_so": "3",  "ma_so_mau": "Blank-01",       "kq": None,    "ghi_chu": ""},
        {"lo_so": "4",  "ma_so_mau": "Spike-01",       "kq": "1.02",  "ghi_chu": ""},
        {"lo_so": "5",  "ma_so_mau": "M03-2026-TFL",  "kq": None,    "ghi_chu": ""},
        {"lo_so": "6",  "ma_so_mau": "M04-2026-TFL",  "kq": "0.015", "ghi_chu": "f=2"},
    ]
}

TEMPLATE_FILE = "filebieumau2.docx"
OUTPUT_FILE   = "TEST_REPORT_OUTPUT.docx"

# ── BƯỚC 1: Điền dữ liệu vào Word template ──────────────────────────
def fill_template(template_path, data, output_path):
    shutil.copy(template_path, output_path)
    doc = Document(output_path)

    # Xóa nội dung thừa, chỉ giữ form page
    body = doc.element.body
    children = list(body)
    KEEP = set(range(171, 203))
    KEEP.add(257)
    for child in [children[i] for i in range(len(children)) if i not in KEEP]:
        try: body.remove(child)
        except: pass

    def fill_cell(cell, text, font_pt=9):
        para = cell.paragraphs[0]
        for run in para.runs:
            run.text = ''
        if text:
            run = para.add_run(str(text))
            run.font.size = Pt(font_pt)

    def replace_in_para(search, replacement):
        for para in doc.paragraphs:
            full = ''.join(r.text for r in para.runs)
            if search in full:
                new = full.replace(search, replacement)
                for r in para.runs[1:]: r.text = ''
                if para.runs: para.runs[0].text = new
                return True
        return False

    # Checkbox kết quả tổng quát
    nd_char   = '☑' if data['check_tat_ca_nd']         else '☐'
    phat_char = '☑' if data['check_co_mau_phat_hien']  else '☐'
    replace_in_para('Các mẫu thử không phát hiện Trifluralin',
                    f'{nd_char} Các mẫu thử không phát hiện Trifluralin')
    replace_in_para('Có mẫu thử phát hiện Trifluralin',
                    f'{phat_char} Có mẫu thử phát hiện Trifluralin')

    # Ngày / người
    replace_in_para('Ngày/Người phân tích',
                    f"{data['ngay_phan_tich']} / {data['nguoi_phan_tich']}")
    replace_in_para('Ngày/Người thẩm tra',
                    f"{data['ngay_tham_tra']} / {data['nguoi_tham_tra']}")

    # Bảng mẫu (table index 2 sau khi xóa)
    tbl = doc.tables[2]
    for i, sample in enumerate(data['samples']):
        ri = i + 1
        if ri >= len(tbl.rows): break
        row = tbl.rows[ri]
        fill_cell(row.cells[0], sample['lo_so'])
        fill_cell(row.cells[1], sample['ma_so_mau'])
        fill_cell(row.cells[2], sample['kq'] if sample['kq'] else '')  # trống = ND
        fill_cell(row.cells[3], sample['ghi_chu'])

    doc.save(output_path)
    print(f"✅ Bước 1 OK — đã lưu: {output_path}")
    return output_path


# ── BƯỚC 2: Upload lên Google Drive → Export PDF ─────────────────────
def upload_and_export_pdf(docx_path, service_account_json_str):
    from google.oauth2 import service_account
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaFileUpload, MediaIoBaseDownload

    creds_data = json.loads(service_account_json_str)
    scopes = ['https://www.googleapis.com/auth/drive']
    creds = service_account.Credentials.from_service_account_info(creds_data, scopes=scopes)
    service = build('drive', 'v3', credentials=creds, cache_discovery=False)

    # Upload .docx → Google Drive (convert to Google Docs)
    print("⏳ Uploading to Google Drive...")
    file_metadata = {
        'name': 'LIMS_Test_Report',
        'mimeType': 'application/vnd.google-apps.document'  # auto-convert
    }
    media = MediaFileUpload(
        docx_path,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    uploaded = service.files().create(
        body=file_metadata,
        media_body=media,
        fields='id,name,webViewLink'
    ).execute()

    file_id = uploaded['id']
    print(f"✅ Bước 2a OK — uploaded: {uploaded['name']} (ID: {file_id})")
    print(f"   Google Docs URL: {uploaded.get('webViewLink', 'N/A')}")

    # Export as PDF
    print("⏳ Exporting as PDF...")
    request = service.files().export_media(fileId=file_id, mimeType='application/pdf')
    pdf_buffer = io.BytesIO()
    downloader = MediaIoBaseDownload(pdf_buffer, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()

    pdf_path = docx_path.replace('.docx', '.pdf')
    with open(pdf_path, 'wb') as f:
        f.write(pdf_buffer.getvalue())
    print(f"✅ Bước 2b OK — PDF saved: {pdf_path} ({len(pdf_buffer.getvalue())} bytes)")

    # Xóa file tạm trên Drive
    service.files().delete(fileId=file_id).execute()
    print(f"🗑️  Đã xóa file tạm trên Drive")

    return pdf_path


# ── MAIN ─────────────────────────────────────────────────────────────
if __name__ == '__main__':
    # Bước 1: Fill template
    output = fill_template(TEMPLATE_FILE, TEST_DATA, OUTPUT_FILE)

    # Bước 2: Upload & PDF
    sa_json = os.environ.get('FIREBASE_SERVICE_ACCOUNT')
    if not sa_json:
        print("\n⚠️  Không tìm thấy FIREBASE_SERVICE_ACCOUNT env variable.")
        print("   Bước 2 bị skip. Chỉ kiểm tra file Word đã điền.")
        print(f"   File output: {OUTPUT_FILE}")
        import subprocess
        subprocess.run(['open', OUTPUT_FILE])
    else:
        try:
            pdf = upload_and_export_pdf(output, sa_json)
            print(f"\n🎉 THÀNH CÔNG! PDF: {pdf}")
            import subprocess
            subprocess.run(['open', pdf])
        except Exception as e:
            print(f"\n❌ Lỗi Drive API: {e}")
            print("   Có thể cần enable Drive API trong Google Cloud Console")
            print(f"   File Word vẫn OK: {OUTPUT_FILE}")
