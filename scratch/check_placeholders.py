import re
from docx import Document

def find_placeholders(docx_path):
    print(f"\n=== Placeholders in {docx_path} ===")
    doc = Document(docx_path)
    pattern = re.compile(r'\{\{[^}]+\}\}')
    
    # Check paragraphs
    for p_idx, para in enumerate(doc.paragraphs):
        matches = pattern.findall(para.text)
        if matches:
            print(f"  Para {p_idx}: {para.text} -> Matches: {matches}")
            
    # Check tables
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                matches = pattern.findall(cell.text)
                if matches:
                    print(f"  Table {t_idx} Row {r_idx} Col {c_idx}: {cell.text} -> Matches: {matches}")

if __name__ == '__main__':
    find_placeholders('filebieumau2_TEMPLATE.docx')
    find_placeholders('filebieumau2_FORM_CLEAN.docx')
    find_placeholders('filebieumau2.docx')
