import os
from docx import Document

files = [
    'filebieumau2.docx',
    'filebieumau2_TEMPLATE.docx',
    'filebieumau2_FORM_CLEAN.docx',
    'filebieumau2_GDOCS_TEMPLATE.docx',
    'filebieumau2_SIMPLE_TEMPLATE.docx'
]

for f in files:
    if os.path.exists(f):
        doc = Document(f)
        print(f"\n=== File: {f} ===")
        print(f"Total tables: {len(doc.tables)}")
        for idx, table in enumerate(doc.tables):
            print(f"  Table {idx}: Rows = {len(table.rows)}, Cols = {len(table.columns)}")
            for r_idx, row in enumerate(table.rows[:2]):
                cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                print(f"    Row {r_idx}: {cells}")
