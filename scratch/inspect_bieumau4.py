import re
import os
from docx import Document

def inspect_docx(docx_path):
    print(f"=== Inspecting {os.path.basename(docx_path)} ===")
    if not os.path.exists(docx_path):
        print(f"Error: {docx_path} does not exist!")
        return

    doc = Document(docx_path)
    
    # 1. Print first 15 paragraphs to see title and metadata
    print("\n--- First Paragraphs ---")
    for i, para in enumerate(doc.paragraphs[:20]):
        text = para.text.strip()
        if text:
            print(f"Para {i}: {text}")
            
    # 2. Print tables info
    print("\n--- Tables Info ---")
    print(f"Total tables: {len(doc.tables)}")
    for t_idx, table in enumerate(doc.tables):
        print(f"Table {t_idx}: {len(table.rows)} rows, max {max(len(row.cells) for row in table.rows)} columns")
        # Print first row of the table to understand headers
        for r_idx in range(min(5, len(table.rows))):
            row = table.rows[r_idx]
            cell_texts = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            print(f"  Row {r_idx}: {cell_texts[:10]}")

    # 3. Find placeholders
    print("\n--- Placeholders found ---")
    pattern = re.compile(r'\{\{[^}]+\}\}')
    
    # Check paragraphs
    for p_idx, para in enumerate(doc.paragraphs):
        matches = pattern.findall(para.text)
        if matches:
            print(f"  Para {p_idx}: {para.text} -> {matches}")
            
    # Check tables
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                matches = pattern.findall(cell.text)
                if matches:
                    print(f"  Table {t_idx} Row {r_idx} Col {c_idx}: {cell.text.strip()} -> {matches}")

if __name__ == '__main__':
    inspect_docx(r'c:\Users\GCMS\Documents\GitHub\lims\FILEBIEUMAUGOC\filebieumau4.docx')
