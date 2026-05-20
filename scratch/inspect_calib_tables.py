import docx

def inspect_tables_3_4(file_path):
    doc = docx.Document(file_path)
    print("--- Table 3 ---")
    for r_idx, row in enumerate(doc.tables[3].rows):
        cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
        print(f"Row {r_idx}: {cells}")

    print("--- Table 4 ---")
    for r_idx, row in enumerate(doc.tables[4].rows):
        cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
        print(f"Row {r_idx}: {cells}")

inspect_tables_3_4("filebieumau2.docx")
