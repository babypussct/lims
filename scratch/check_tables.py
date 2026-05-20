import os
from docx import Document

def dump_tables(docx_path):
    print(f"=== Inspecting document: {docx_path} ===")
    if not os.path.exists(docx_path):
        print("File does not exist!")
        return
    
    doc = Document(docx_path)
    print(f"Total tables: {len(doc.tables)}")
    for idx, table in enumerate(doc.tables):
        print(f"\nTable {idx}: Rows = {len(table.rows)}, Cols = {len(table.columns)}")
        # Print top rows
        for r_idx, row in enumerate(table.rows[:10]):
            cells_text = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            print(f"  Row {r_idx}: {cells_text}")
        if len(table.rows) > 10:
            print(f"  ... and {len(table.rows) - 10} more rows")

if __name__ == '__main__':
    dump_tables('filebieumau2.docx')
