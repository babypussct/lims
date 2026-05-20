import docx

doc = docx.Document("filebieumau2.docx")
print("=== Table 3 ===")
for r_idx, row in enumerate(doc.tables[3].rows):
    row_text = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
    print(f"Row {r_idx}: {row_text}")

print("=== Table 4 ===")
for r_idx, row in enumerate(doc.tables[4].rows):
    row_text = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
    print(f"Row {r_idx}: {row_text}")
