from docx import Document

doc = Document('filebieumau2.docx')
body = doc.element.body
children = list(body)

# The report form starts around index 170. Let's inspect the tags and text/content of children from 170 to the end.
for idx in range(170, len(children)):
    child = children[idx]
    tag = child.tag.split('}')[-1]
    
    if tag == 'p':
        # Find paragraph text
        # Find paragraph in doc.paragraphs
        p_text = ""
        for p in doc.paragraphs:
            if p._element is child:
                p_text = p.text.strip()
                break
        if p_text:
            print(f"Index {idx} [P]: {p_text}")
    elif tag == 'tbl':
        # Find table in doc.tables
        table_idx = -1
        for t_i, t in enumerate(doc.tables):
            if t._element is child:
                table_idx = t_i
                break
        print(f"Index {idx} [Table {table_idx}]: Rows={len(doc.tables[table_idx].rows)}, Cols={len(doc.tables[table_idx].columns)}")
        for r_i, row in enumerate(doc.tables[table_idx].rows):
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            print(f"  Row {r_i}: {cells}")
