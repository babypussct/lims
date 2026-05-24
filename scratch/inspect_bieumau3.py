import docx

doc = docx.Document("filebieumau3.docx")
print("Total paragraphs:", len(doc.paragraphs))
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"P{i}: {para.text.strip()}")

print("\nTotal tables:", len(doc.tables))
for i, table in enumerate(doc.tables):
    rows = len(table.rows)
    cols = len(table.columns)
    first_cell_text = ""
    if rows > 0 and cols > 0:
        first_cell_text = table.cell(0, 0).text.strip().replace('\n', ' ')
    print(f"Table {i}: {rows} rows x {cols} cols, first cell text: '{first_cell_text}'")
