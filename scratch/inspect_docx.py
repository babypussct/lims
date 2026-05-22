import os
from docx import Document

doc_path = r"c:\Users\GCMS\Documents\GitHub\lims\FORM_GOC_TRIFLURALIN_9_3.docx"
if not os.path.exists(doc_path):
    print(f"File not found: {doc_path}")
    exit(1)

doc = Document(doc_path)
print(f"Total tables: {len(doc.tables)}")
for i, table in enumerate(doc.tables):
    rows = len(table.rows)
    cols = len(table.columns) if rows > 0 else 0
    first_cell_text = ""
    if rows > 0 and cols > 0:
        first_cell_text = table.rows[0].cells[0].text.strip().replace('\n', ' ')
    print(f"Table {i}: {rows} rows x {cols} cols. First cell: '{first_cell_text}'")
    # Print the last row's first cell to see if it contains R^2 or signatures
    if rows > 0:
        last_cell_text = table.rows[-1].cells[0].text.strip().replace('\n', ' ')
        print(f"  Last cell: '{last_cell_text}'")
