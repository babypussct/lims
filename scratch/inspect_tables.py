import docx

doc = docx.Document("filebieumau2.docx")
print("Total tables:", len(doc.tables))
for i, table in enumerate(doc.tables):
    rows = len(table.rows)
    cols = len(table.columns)
    first_cell_text = ""
    if rows > 0 and cols > 0:
        first_cell_text = table.cell(0, 0).text.strip()
    print(f"Table {i}: {rows} rows x {cols} cols, first cell text: '{first_cell_text}'")
    # Print first row cells
    cells = [table.cell(0, c).text.strip().replace('\n', ' ') for c in range(cols)]
    print(f"  Header: {cells}")
