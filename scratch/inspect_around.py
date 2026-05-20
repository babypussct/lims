import docx

def inspect_around(file_path):
    doc = docx.Document(file_path)
    for i in range(175, 215):
        if i < len(doc.paragraphs):
            print(f"Para {i}: {doc.paragraphs[i].text.strip()}")
            
    print("\n--- Tables and their relative order ---")
    body = doc.element.body
    children = list(body)
    for i, child in enumerate(children):
        if child.tag.endswith('tbl'):
            # find table index in doc.tables
            for t_idx, t in enumerate(doc.tables):
                if t._tbl == child:
                    print(f"Child {i} is Table {t_idx}")
        elif child.tag.endswith('p'):
            text = child.text if child.text else ""
            if any(k in text for k in ["Đường chuẩn", "Mẫu phân tích", "5.", "6.", "7."]):
                print(f"Child {i} is Paragraph: {text.strip()}")

inspect_around("filebieumau2.docx")
